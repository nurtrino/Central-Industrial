"""
Monkey Read Monkey Do — notes generation from a transcript.

Sends ONLY the text transcript (never audio) to the Claude API, prompted with the
user's house style guide, and renders the result to (a) markdown/HTML for the
browser window and (b) a Word .docx in the work folder.
"""

import os
import re

STYLE_GUIDE_PATH = os.path.join(os.path.dirname(__file__),
                                "transcipt and summarize notes.md")

# Opus for the high-stakes DD notes; override with NOTEMAX_MODEL in .env.
DEFAULT_MODEL = os.environ.get("NOTEMAX_MODEL", "claude-opus-4-8")


def load_style_guide() -> str:
    with open(STYLE_GUIDE_PATH, "r", encoding="utf-8") as f:
        return f.read()


def _stream_text(client, model, system, content, log, control=None):
    """Stream a Claude message, accumulating text. Honors stop (break on control.stopped)."""
    out = []
    with client.messages.stream(
        model=model, max_tokens=16000, system=system,
        messages=[{"role": "user", "content": content}],
    ) as stream:
        for chunk in stream.text_stream:
            out.append(chunk)
            if control is not None and control.stopped:
                break                          # stop -> keep partial text
    return "".join(out)


def generate_notes(transcript: str, api_key: str, model: str = None,
                   extra_context: str = "", image_paths=None, log=print,
                   control=None) -> str:
    """Call Claude with the style guide + transcript; return markdown notes.

    image_paths: optional list of supplied slide/screenshot files. Per the style
    guide step 3, Claude reviews them and embeds image placeholders in the notes.
    """
    import anthropic

    model = model or DEFAULT_MODEL
    style = load_style_guide()
    client = anthropic.Anthropic(api_key=api_key)

    system = (
        "You are an expert note-taker for meetings, calls, and presentations. You convert a "
        "diarized meeting transcript into written notes that follow the house style "
        "guide below EXACTLY. The notes are the deliverable — they must be accurate, "
        "complete on key points, and free of outside knowledge.\n\n"
        "Output ONLY the notes themselves in GitHub-flavored Markdown. Do not add a "
        "preamble, sign-off, or commentary about your process. Use Markdown pipe "
        "tables for any genuine tabular data. Use **bold** (never varied font sizes) "
        "and bullet points per the guide.\n\n"
        "==== HOUSE STYLE GUIDE ====\n" + style
    )

    # Build the user content: transcript (+ any images as exhibits to place).
    content = []
    if image_paths:
        content.append({
            "type": "text",
            "text": (f"The following {len(image_paths)} image(s) were supplied as "
                     "slides/screenshots/exhibits. Per the style guide, review each and "
                     "place it INTO the notes at the most contextually relevant point — "
                     "near the discussion it illustrates. Insert a placeholder of the "
                     "form `[[EXHIBIT n]]` (1-based, matching the order below) ON ITS "
                     "OWN LINE (blank line before and after) at that point. You MUST "
                     "place every exhibit exactly once; if you cannot tell where one "
                     "belongs, put its placeholder at the very end of the notes. If an "
                     "image contains a genuine DATA TABLE, also reproduce that table as "
                     "a Markdown pipe table near the placeholder (do not restate every "
                     "cell as prose)."),
        })
        for i, p in enumerate(image_paths, 1):
            content.append({"type": "text", "text": f"--- EXHIBIT {i}: {os.path.basename(p)} ---"})
            content.append(_image_block(p))

    if extra_context:
        content.append({"type": "text",
                        "text": f"Additional supplied context (already extracted to "
                                f"text):\n\n{extra_context}"})

    content.append({"type": "text",
                    "text": "==== TRANSCRIPT ====\n\n" + transcript})

    log(f"  generating notes with {model} (transcript ~{len(transcript)} chars)...")
    return _stream_text(client, model, system, content, log, control).strip()


EDITOR_GUIDE_PATH = os.path.join(os.path.dirname(__file__), "notes_editor_review.md")


def load_editor_guide() -> str:
    with open(EDITOR_GUIDE_PATH, "r", encoding="utf-8") as f:
        return f.read()


def review_notes(draft_md: str, transcript: str, api_key: str, model: str = None,
                 extra_context: str = "", image_paths=None, log=print, control=None):
    """Second-pass senior editorial review. Returns (final_md, changelog).

    Sends the spec, transcript (ground truth), supplied docs/images and the draft
    to a fresh reviewer per notes_editor_review.md. Translates our [[EXHIBIT n]]
    tokens to the editor's {{IMG:imgN}} contract and back.
    """
    import anthropic

    model = model or DEFAULT_MODEL
    client = anthropic.Anthropic(api_key=api_key)
    system = load_editor_guide()

    # [[EXHIBIT n]] -> {{IMG:imgN}} for the editor.
    draft_tok = _EXHIBIT_RE.sub(lambda m: "{{IMG:img%s}}" % m.group(1), draft_md)

    content = [
        {"type": "text", "text": "1. THE SPEC (style tips the notes must follow):\n\n"
                                 + load_style_guide()},
        {"type": "text", "text": "2. THE TRANSCRIPT (ground truth):\n\n" + transcript},
    ]
    if extra_context:
        content.append({"type": "text",
                        "text": "3. SUPPLIED DOCUMENTS (extracted text):\n\n" + extra_context})
    if image_paths:
        content.append({"type": "text", "text": "4. SUPPLIED IMAGES (each with its key):"})
        for i, p in enumerate(image_paths, 1):
            content.append({"type": "text", "text": f"img{i}: {os.path.basename(p)}"})
            content.append(_image_block(p))
    content.append({"type": "text",
                    "text": "5. THE DRAFT NOTES (markdown, with {{IMG:imgK}} tokens):\n\n"
                            + draft_tok})

    log(f"  editorial review pass with {model}...")
    raw = _stream_text(client, model, system, content, log, control)

    # Split off the changelog.
    final, changelog = raw, ""
    if "<<<EDITOR_CHANGELOG>>>" in raw:
        final, _, changelog = raw.partition("<<<EDITOR_CHANGELOG>>>")

    # {{IMG:imgN}} -> [[EXHIBIT n]] back for our renderer.
    final = re.sub(r"\{\{\s*IMG:img(\d+)\s*\}\}",
                   lambda m: "[[EXHIBIT %s]]" % m.group(1), final, flags=re.IGNORECASE)
    return final.strip(), changelog.strip()


_EXHIBIT_RE = re.compile(r"\[\[\s*EXHIBIT\s+(\d+)\s*\]\]", re.IGNORECASE)


def normalize_exhibits(md: str, num_images: int) -> str:
    """Guarantee every supplied image appears exactly once, each on its own line.

    - Forces each `[[EXHIBIT n]]` onto its own line (so the docx/HTML renderers
      embed it reliably even if the model wrote it inline).
    - De-duplicates repeated placeholders for the same exhibit.
    - Appends any exhibit the model failed to place at the end (per the guide:
      "If it cannot be determined where to put any or all of the images, put them
      at the end").
    """
    if num_images <= 0:
        return md

    seen = set()

    def repl(m):
        n = int(m.group(1))
        if n < 1 or n > num_images or n in seen:
            return ""               # drop out-of-range or duplicate placeholders
        seen.add(n)
        return f"\n\n[[EXHIBIT {n}]]\n\n"

    md = _EXHIBIT_RE.sub(repl, md)

    missing = [n for n in range(1, num_images + 1) if n not in seen]
    if missing:
        md = md.rstrip() + "\n\n**Exhibits**\n" + \
            "".join(f"\n[[EXHIBIT {n}]]\n" for n in missing)

    # collapse the 3+ blank lines the insertions may create
    md = re.sub(r"\n{3,}", "\n\n", md)
    return md.strip()


def _image_block(path: str):
    import base64
    ext = os.path.splitext(path)[1].lower().lstrip(".")
    media = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
             "gif": "image/gif", "webp": "image/webp"}.get(ext, "image/png")
    with open(path, "rb") as f:
        data = base64.standard_b64encode(f.read()).decode("ascii")
    return {"type": "image",
            "source": {"type": "base64", "media_type": media, "data": data}}


# --------------------------------------------------------------------------- #
# Markdown -> Word .docx  (lightweight; supports the style guide's limited set:
# headings, bold, bullets, pipe tables, exhibit images, paragraphs)
# --------------------------------------------------------------------------- #
def markdown_to_docx(md: str, out_path: str, title: str = None, image_map=None):
    """Render markdown notes to a .docx. image_map maps '[[EXHIBIT n]]' -> path."""
    from docx import Document
    from docx.shared import Pt, Inches

    image_map = image_map or {}
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        h = doc.add_paragraph()
        run = h.add_run(title)
        run.bold = True
        run.font.size = Pt(14)

    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Exhibit placeholder on its own line -> embed image.
        ex = re.fullmatch(r"\s*\[\[EXHIBIT\s+(\d+)\]\]\s*", line, re.IGNORECASE)
        if ex and image_map.get(int(ex.group(1))):
            try:
                doc.add_picture(image_map[int(ex.group(1))], width=Inches(6.0))
            except Exception:
                doc.add_paragraph(f"[Exhibit {ex.group(1)} — image could not be embedded]")
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        # Pipe table: a header row followed by a |---| separator.
        if line.lstrip().startswith("|") and i + 1 < len(lines) and \
                re.match(r"\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]) and "-" in lines[i + 1]:
            table_lines = [line]
            j = i + 1
            sep = lines[j]
            j += 1
            while j < len(lines) and lines[j].lstrip().startswith("|"):
                table_lines.append(lines[j])
                j += 1
            _add_table(doc, [table_lines[0]] + table_lines[1:])
            i = j
            continue

        # Heading.
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            p = doc.add_paragraph()
            run = p.add_run(m.group(2).strip())
            run.bold = True            # guide: no varied font sizes, bold only
            i += 1
            continue

        # Bullet.
        mb = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if mb:
            indent = len(mb.group(1)) // 2
            p = doc.add_paragraph(style="List Bullet")
            if indent:
                p.paragraph_format.left_indent = Inches(0.25 * indent)
            _add_inline(p, mb.group(2))
            i += 1
            continue

        # Plain paragraph.
        p = doc.add_paragraph()
        _add_inline(p, line)
        i += 1

    doc.save(out_path)
    return out_path


def _add_table(doc, rows):
    from docx.shared import Pt

    def cells(row):
        return [c.strip() for c in row.strip().strip("|").split("|")]

    header = cells(rows[0])
    body = [cells(r) for r in rows[2:]] if len(rows) > 2 else []
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for k, text in enumerate(header):
        run = table.rows[0].cells[k].paragraphs[0].add_run(text)
        run.bold = True
    for row in body:
        cell_row = table.add_row().cells
        for k in range(len(header)):
            txt = row[k] if k < len(row) else ""
            _add_inline(cell_row[k].paragraphs[0], txt)


def _add_inline(paragraph, text):
    """Handle **bold** runs within a line."""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
